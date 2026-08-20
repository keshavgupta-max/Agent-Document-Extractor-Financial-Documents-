"""DOCX Document Parser using python-docx."""

import io
import docx
from logger import logger
from tools.parser.exceptions import CorruptedDocument, ParserExecutionError
from tools.parser.models import PageContent, ParsedDocument, TableContent


class DOCXParser:
    """Parser dedicated to extracting text, tables, and metadata from Microsoft Word (.docx) files."""

    def parse(
        self,
        file_bytes: bytes,
        document_id: str,
        storage_path: str,
        mime_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ) -> ParsedDocument:
        """Extracts text paragraphs, tables, and core document metadata from DOCX bytes.

        Raises:
            CorruptedDocument: If the DOCX stream cannot be opened.
            ParserExecutionError: If an error occurs during extraction.
        """
        try:
            doc_stream = io.BytesIO(file_bytes)
            doc = docx.Document(doc_stream)
        except Exception as exc:
            error_msg = f"Failed to open DOCX document (corrupted or invalid): {str(exc)}"
            logger.error(error_msg)
            raise CorruptedDocument(error_msg) from exc

        try:
            # Extract paragraph text
            paragraphs_text = [
                p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()
            ]
            full_text = "\n\n".join(paragraphs_text)

            # Treat entire document as 1 logical section/page unless page breaks are handled
            pages = [PageContent(page_number=1, text=full_text)]

            # Extract tables
            tables = []
            for idx, table in enumerate(doc.tables):
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_data)

                if table_rows:
                    headers = table_rows[0]
                    rows = table_rows[1:] if len(table_rows) > 1 else []
                    tables.append(
                        TableContent(
                            table_index=idx,
                            page_number=1,
                            headers=headers,
                            rows=rows,
                        )
                    )

            # Metadata extraction
            core_properties = doc.core_properties
            metadata = {
                "author": core_properties.author or "",
                "title": core_properties.title or "",
                "subject": core_properties.subject or "",
                "created": str(core_properties.created) if core_properties.created else "",
                "modified": str(core_properties.modified) if core_properties.modified else "",
            }

            return ParsedDocument(
                document_id=document_id,
                storage_path=storage_path,
                file_extension=".docx",
                mime_type=mime_type,
                page_count=1,
                pages=pages,
                tables=tables,
                metadata=metadata,
                parsing_status="SUCCESS",
            )

        except Exception as exc:
            error_msg = f"Error extracting DOCX content for document {document_id}: {str(exc)}"
            logger.error(error_msg, exc_info=True)
            raise ParserExecutionError(error_msg) from exc